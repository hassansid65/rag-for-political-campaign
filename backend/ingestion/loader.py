"""
Multi-format document loader: PDF / DOCX / TXT / MD (+ HTML, CSV).

Design notes
------------
* Every loader returns `LoadedDocument` with page-aligned `blocks` so the chunker
  can keep page numbers and heading breadcrumbs on each chunk. Losing page
  provenance here makes citations useless downstream.
* `unstructured` is supported but optional — it is a heavy dependency and its
  PDF partitioner is an order of magnitude slower than PyMuPDF. Set
  `USE_UNSTRUCTURED=1` to prefer it (better at tables and multi-column layouts);
  otherwise we use PyMuPDF and reconstruct structure from font sizes.
"""

from __future__ import annotations

import csv
import io
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Optional

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".txt",
    ".md",
    ".markdown",
    ".html",
    ".htm",
    ".csv",
}

USE_UNSTRUCTURED = os.getenv("USE_UNSTRUCTURED", "0").lower() in {"1", "true", "yes"}


@dataclass
class Block:
    """A structural unit of text with its provenance."""

    text: str
    page: Optional[int] = None
    kind: str = "paragraph"          # paragraph | heading | list_item | table | caption
    level: int = 0                   # heading depth (1 = top)


@dataclass
class LoadedDocument:
    source: str
    blocks: list[Block]
    pages: Optional[int] = None
    detected_language: str = "en"
    warnings: list[str] = field(default_factory=list)
    raw_text: str = ""

    @property
    def char_count(self) -> int:
        return len(self.raw_text)


class UnsupportedFileType(ValueError):
    pass


# --------------------------------------------------------------------- helpers
_LIGATURES = {
    "ﬀ": "ff", "ﬁ": "fi", "ﬂ": "fl", "ﬃ": "ffi", "ﬄ": "ffl",
    "“": '"', "”": '"', "‘": "'", "’": "'",
    "–": "-", "—": "-", " ": " ", "​": "",
}


def clean_text(text: str) -> str:
    """Normalize PDF extraction artefacts without destroying layout signal."""
    for bad, good in _LIGATURES.items():
        text = text.replace(bad, good)
    # De-hyphenate words split across line breaks: "Vijaya-\nwada" -> "Vijayawada"
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    # Collapse 3+ newlines to a paragraph break, and trailing spaces.
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


_TELUGU = re.compile(r"[ఀ-౿]")
_DEVANAGARI = re.compile(r"[ऀ-ॿ]")
_TAMIL = re.compile(r"[஀-௿]")


def detect_language(text: str) -> str:
    """Script-based language hint — enough to pick a TTS voice and STT locale."""
    sample = text[:4000]
    if _TELUGU.search(sample):
        return "te"
    if _DEVANAGARI.search(sample):
        return "hi"
    if _TAMIL.search(sample):
        return "ta"
    return "en"


# ------------------------------------------------------------------ PDF loader
def _load_pdf_pymupdf(path: Path) -> LoadedDocument:
    import fitz  # PyMuPDF

    blocks: list[Block] = []
    warnings: list[str] = []
    doc = fitz.open(path)
    try:
        # Body font size is the modal size across the doc; anything meaningfully
        # larger and short is a heading. This recovers structure that plain
        # get_text() throws away, and heading breadcrumbs measurably improve
        # retrieval precision on manifesto-style documents.
        sizes: dict[int, int] = {}
        for page in doc:
            for blk in page.get_text("dict").get("blocks", []):
                for line in blk.get("lines", []):
                    for span in line.get("spans", []):
                        if span.get("text", "").strip():
                            key = int(round(span.get("size", 10)))
                            sizes[key] = sizes.get(key, 0) + len(span["text"])
        body_size = max(sizes, key=lambda k: sizes[k]) if sizes else 10

        for page_no, page in enumerate(doc, start=1):
            page_dict = page.get_text("dict")
            for blk in page_dict.get("blocks", []):
                if blk.get("type") != 0:  # 0 = text
                    continue
                lines: list[str] = []
                max_size = 0.0
                bold = False
                for line in blk.get("lines", []):
                    parts = [s.get("text", "") for s in line.get("spans", [])]
                    for span in line.get("spans", []):
                        max_size = max(max_size, float(span.get("size", 0)))
                        if "bold" in str(span.get("font", "")).lower():
                            bold = True
                    joined = "".join(parts).strip()
                    if joined:
                        lines.append(joined)
                if not lines:
                    continue
                text = clean_text("\n".join(lines))
                if not text:
                    continue

                is_short = len(text) <= 120 and text.count("\n") <= 1
                if is_short and (max_size >= body_size + 1.5 or (bold and max_size >= body_size)):
                    level = 1 if max_size >= body_size + 4 else 2
                    blocks.append(Block(text=text, page=page_no, kind="heading", level=level))
                elif re.match(r"^\s*(?:[-•*•]|\d+[.)])\s+", lines[0]):
                    blocks.append(Block(text=text, page=page_no, kind="list_item"))
                else:
                    blocks.append(Block(text=text, page=page_no, kind="paragraph"))

        # Extract tables separately when PyMuPDF supports it — table rows carry
        # the scheme amounts and beneficiary counts voters actually ask about.
        for page_no, page in enumerate(doc, start=1):
            try:
                finder = page.find_tables()
            except Exception:  # older PyMuPDF, or no table support
                break
            for table in getattr(finder, "tables", []) or []:
                try:
                    rows = table.extract()
                except Exception:
                    continue
                rendered = _render_table(rows)
                if rendered:
                    blocks.append(Block(text=rendered, page=page_no, kind="table"))

        page_count = doc.page_count
    finally:
        doc.close()

    if not blocks:
        warnings.append(
            "No extractable text found — the PDF is likely scanned. "
            "Run it through OCR (e.g. ocrmypdf) before ingesting."
        )

    raw = "\n\n".join(b.text for b in blocks)
    return LoadedDocument(
        source=path.name,
        blocks=blocks,
        pages=page_count,
        detected_language=detect_language(raw),
        warnings=warnings,
        raw_text=raw,
    )


def _render_table(rows: Iterable[Iterable[Optional[str]]]) -> str:
    """Markdown-ish table rendering — keeps header/value association for the LLM."""
    cleaned: list[list[str]] = []
    for row in rows:
        cells = [(c or "").strip().replace("\n", " ") for c in row]
        if any(cells):
            cleaned.append(cells)
    if len(cleaned) < 2:
        return ""
    header, *body = cleaned
    out = [" | ".join(header), " | ".join("---" for _ in header)]
    out.extend(" | ".join(r) for r in body)
    return "\n".join(out)


def _load_with_unstructured(path: Path) -> LoadedDocument:
    from unstructured.partition.auto import partition

    elements = partition(filename=str(path))
    kind_map = {
        "Title": ("heading", 1),
        "Header": ("heading", 2),
        "ListItem": ("list_item", 0),
        "Table": ("table", 0),
        "FigureCaption": ("caption", 0),
    }
    blocks: list[Block] = []
    for el in elements:
        text = clean_text(str(el))
        if not text:
            continue
        kind, level = kind_map.get(type(el).__name__, ("paragraph", 0))
        page = getattr(getattr(el, "metadata", None), "page_number", None)
        blocks.append(Block(text=text, page=page, kind=kind, level=level))

    raw = "\n\n".join(b.text for b in blocks)
    pages = max((b.page or 0) for b in blocks) or None
    return LoadedDocument(
        source=path.name,
        blocks=blocks,
        pages=pages,
        detected_language=detect_language(raw),
        raw_text=raw,
    )


# ----------------------------------------------------------------- DOCX loader
def _load_docx(path: Path) -> LoadedDocument:
    import docx  # python-docx

    document = docx.Document(str(path))
    blocks: list[Block] = []

    for para in document.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            match = re.search(r"(\d+)", style)
            level = int(match.group(1)) if match else 1
            blocks.append(Block(text=text, kind="heading", level=level))
        elif style.startswith("list") or style.startswith("bullet"):
            blocks.append(Block(text=text, kind="list_item"))
        elif style in {"title", "subtitle"}:
            blocks.append(Block(text=text, kind="heading", level=1))
        else:
            blocks.append(Block(text=text, kind="paragraph"))

    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        rendered = _render_table(rows)
        if rendered:
            blocks.append(Block(text=rendered, kind="table"))

    raw = "\n\n".join(b.text for b in blocks)
    return LoadedDocument(
        source=path.name,
        blocks=blocks,
        detected_language=detect_language(raw),
        raw_text=raw,
    )


# ------------------------------------------------------- TXT / MD / HTML / CSV
_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_MD_SETEXT_H1 = re.compile(r"^={3,}\s*$")
_MD_SETEXT_H2 = re.compile(r"^-{3,}\s*$")


def _load_markdown(path: Path, text: Optional[str] = None) -> LoadedDocument:
    content = text if text is not None else path.read_text(encoding="utf-8", errors="replace")
    lines = content.splitlines()
    blocks: list[Block] = []
    buffer: list[str] = []
    in_fence = False
    in_table = False
    table_buf: list[str] = []

    def flush_para() -> None:
        if buffer:
            cleaned = clean_text("\n".join(buffer))
            if cleaned:
                kind = (
                    "list_item"
                    if re.match(r"^\s*(?:[-*+•]|\d+[.)])\s+", buffer[0])
                    else "paragraph"
                )
                blocks.append(Block(text=cleaned, kind=kind))
            buffer.clear()

    def flush_table() -> None:
        nonlocal in_table
        if table_buf:
            blocks.append(Block(text="\n".join(table_buf), kind="table"))
            table_buf.clear()
        in_table = False

    for i, line in enumerate(lines):
        if line.strip().startswith("```"):
            in_fence = not in_fence
            buffer.append(line)
            continue
        if in_fence:
            buffer.append(line)
            continue

        # Markdown pipe table
        if "|" in line and line.strip().startswith("|"):
            flush_para()
            in_table = True
            table_buf.append(line.strip())
            continue
        if in_table:
            flush_table()

        heading = _MD_HEADING.match(line)
        if heading:
            flush_para()
            blocks.append(
                Block(
                    text=clean_text(heading.group(2)),
                    kind="heading",
                    level=len(heading.group(1)),
                )
            )
            continue

        # Setext headings: text underlined with === or ---
        nxt = lines[i + 1] if i + 1 < len(lines) else ""
        if line.strip() and _MD_SETEXT_H1.match(nxt):
            flush_para()
            blocks.append(Block(text=clean_text(line), kind="heading", level=1))
            continue
        if line.strip() and _MD_SETEXT_H2.match(nxt) and not line.strip().startswith("-"):
            flush_para()
            blocks.append(Block(text=clean_text(line), kind="heading", level=2))
            continue
        if _MD_SETEXT_H1.match(line) or (_MD_SETEXT_H2.match(line) and blocks and blocks[-1].kind == "heading"):
            continue

        if not line.strip():
            flush_para()
        else:
            buffer.append(line)

    flush_table()
    flush_para()

    raw = "\n\n".join(b.text for b in blocks)
    return LoadedDocument(
        source=path.name,
        blocks=blocks,
        detected_language=detect_language(raw),
        raw_text=raw,
    )


def _load_txt(path: Path) -> LoadedDocument:
    content = path.read_text(encoding="utf-8", errors="replace")
    # Plain text often *is* markdown-shaped (## headings, bullet lists). Reusing
    # the markdown loader recovers that structure for free.
    if re.search(r"^#{1,6}\s+\S", content, re.MULTILINE):
        return _load_markdown(path, text=content)

    blocks: list[Block] = []
    for para in re.split(r"\n\s*\n", content):
        cleaned = clean_text(para)
        if not cleaned:
            continue
        lines = cleaned.splitlines()
        # An ALL-CAPS or short colon-terminated standalone line reads as a heading.
        if len(lines) == 1 and len(cleaned) <= 90 and (
            cleaned.isupper() or cleaned.endswith(":")
        ):
            blocks.append(Block(text=cleaned.rstrip(":"), kind="heading", level=2))
        else:
            blocks.append(Block(text=cleaned, kind="paragraph"))

    raw = "\n\n".join(b.text for b in blocks)
    return LoadedDocument(
        source=path.name,
        blocks=blocks,
        detected_language=detect_language(raw),
        raw_text=raw,
    )


def _load_html(path: Path) -> LoadedDocument:
    from html.parser import HTMLParser

    class Extractor(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self.blocks: list[Block] = []
            self._buf: list[str] = []
            self._tag: str = "p"
            self._skip = 0

        def handle_starttag(self, tag: str, attrs: list[tuple[str, Optional[str]]]) -> None:
            if tag in {"script", "style"}:
                self._skip += 1
            elif tag in {"p", "li", "div", "td", "h1", "h2", "h3", "h4", "h5", "h6", "br"}:
                self._flush()
                self._tag = tag

        def handle_endtag(self, tag: str) -> None:
            if tag in {"script", "style"}:
                self._skip = max(0, self._skip - 1)
            elif tag in {"p", "li", "div", "td", "h1", "h2", "h3", "h4", "h5", "h6"}:
                self._flush()

        def handle_data(self, data: str) -> None:
            if not self._skip:
                self._buf.append(data)

        def _flush(self) -> None:
            text = clean_text("".join(self._buf))
            self._buf.clear()
            if not text:
                return
            if self._tag and self._tag[0] == "h" and self._tag[1:].isdigit():
                self.blocks.append(Block(text=text, kind="heading", level=int(self._tag[1:])))
            elif self._tag == "li":
                self.blocks.append(Block(text=text, kind="list_item"))
            else:
                self.blocks.append(Block(text=text, kind="paragraph"))

    parser = Extractor()
    parser.feed(path.read_text(encoding="utf-8", errors="replace"))
    parser._flush()

    raw = "\n\n".join(b.text for b in parser.blocks)
    return LoadedDocument(
        source=path.name,
        blocks=parser.blocks,
        detected_language=detect_language(raw),
        raw_text=raw,
    )


def _load_csv(path: Path) -> LoadedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = [r for r in reader if any(c.strip() for c in r)]
    if not rows:
        return LoadedDocument(source=path.name, blocks=[], raw_text="")

    header, *body = rows
    blocks: list[Block] = []
    # One block per row, rendered as "Column: value" pairs. Row-per-chunk keeps
    # each FAQ / district record independently retrievable.
    for row in body:
        pairs = [
            f"{h.strip()}: {v.strip()}"
            for h, v in zip(header, row)
            if v and v.strip()
        ]
        if pairs:
            blocks.append(Block(text="\n".join(pairs), kind="paragraph"))

    raw = "\n\n".join(b.text for b in blocks)
    return LoadedDocument(
        source=path.name,
        blocks=blocks,
        detected_language=detect_language(raw),
        raw_text=raw,
    )


# ------------------------------------------------------------------- dispatch
def load_document(path: str | Path) -> LoadedDocument:
    """Load any supported file into structure-preserving blocks."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(
            f"{ext or '(no extension)'} is not supported. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if USE_UNSTRUCTURED:
        try:
            doc = _load_with_unstructured(path)
            if doc.blocks:
                return doc
            logger.warning("unstructured returned no blocks for %s; falling back", path.name)
        except Exception as exc:  # noqa: BLE001 — any partitioner failure falls back
            logger.warning("unstructured failed for %s (%s); falling back", path.name, exc)

    if ext == ".pdf":
        return _load_pdf_pymupdf(path)
    if ext in {".docx", ".doc"}:
        return _load_docx(path)
    if ext in {".md", ".markdown"}:
        return _load_markdown(path)
    if ext in {".html", ".htm"}:
        return _load_html(path)
    if ext == ".csv":
        return _load_csv(path)
    return _load_txt(path)
